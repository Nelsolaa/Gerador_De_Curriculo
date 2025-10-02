import os
import json
import logging
import io 

from flask import Flask, render_template, request, jsonify, Response
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pypdf import PdfReader

# --- 1. CONFIGURAÇÃO INICIAL ---
load_dotenv()

# Configuração de Logs
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Constantes e configuração do Flask
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

# O nome das pastas de templates e static é ajustado para funcionar na Vercel
import os
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
app = Flask(__name__,
            template_folder=os.path.join(BASE_DIR, "templates"),
            static_folder=os.path.join(BASE_DIR, "static"))

app.config['SECRET_KEY'] = os.getenv('SEGREDO_FLASK')


# --- 2. FUNÇÕES AUXILIARES ---
def allowed_file(filename):
    """Verifica se a extensão do arquivo é permitida."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def ler_arquivo_de_stream(file_stream, filename):
    """Lê o conteúdo de um arquivo diretamente de um stream em memória."""
    logging.info(f"Iniciando a leitura do arquivo em memória: {filename}")
    texto_completo = ""
    try:
        if filename.lower().endswith(".pdf"):
            leitor_pdf = PdfReader(file_stream)
            for pagina in leitor_pdf.pages:
                texto_completo += pagina.extract_text() or ""
        elif filename.lower().endswith(".docx"):
            doc = Document(file_stream)
            for paragrafo in doc.paragraphs:
                texto_completo += paragrafo.text + "\n"
        logging.info("Arquivo lido com sucesso da memória.")
        return texto_completo, True
    except Exception as e:
        logging.error(f"Exceção ao ler o arquivo do stream: {e}", exc_info=True)
        return f"Erro ao ler o arquivo: {e}", False

def criar_curriculo_docx_em_memoria(data):
    """Gera um documento .docx em memória e retorna o buffer de bytes e o nome do arquivo."""
    logging.info("Iniciando a criação do documento DOCX em memória.")
    try:
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        contact_info_data = data.get('contact_info') or {}

        doc.add_paragraph(data.get('name') or 'Nome não encontrado', style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(data.get('role') or 'Cargo não encontrado', style='Subtitle').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        contact_line = " | ".join(filter(None, [
            contact_info_data.get('location'),
            contact_info_data.get('phone'),
            contact_info_data.get('email')
        ]))
        if contact_line:
            doc.add_paragraph(contact_line).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        links_line = " | ".join(filter(None, [
            f"LinkedIn: {contact_info_data.get('linkedin')}" if contact_info_data.get('linkedin') else None,
            f"GitHub: {contact_info_data.get('github')}" if contact_info_data.get('github') else None
        ]))
        if links_line:
            doc.add_paragraph(links_line).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph('_' * 80).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if data.get('profile'):
            doc.add_heading('Profile', level=1)
            p = doc.add_paragraph(data.get('profile') or '')
            p.paragraph_format.space_after = Pt(12)

        if data.get('professional_experience'):
            doc.add_heading('Professional Experience', level=1)
            for exp in data.get('professional_experience', []):
                p = doc.add_paragraph()
                p.add_run((exp.get('title') or '') + '\n').bold = True
                p.add_run(f"{(exp.get('company') or '')}, {(exp.get('location') or '')}\n")
                p.add_run((exp.get('dates') or '') + '\n')
                description_points = (exp.get('description') or '').split('\n')
                formatted_description = "".join([f"- {point.strip()}\n" for point in description_points if point.strip()])
                p.add_run(formatted_description)
                p.paragraph_format.space_after = Pt(12)

        if data.get('education'):
            doc.add_heading('Education', level=1)
            for edu in data.get('education', []):
                p = doc.add_paragraph()
                p.add_run((edu.get('degree') or '') + '\n').bold = True
                p.add_run((edu.get('institution') or '') + '\n')
                p.add_run((edu.get('dates') or '') + '\n')
                if edu.get('description'):
                    p.add_run(f"- {(edu.get('description') or '')}")
                p.paragraph_format.space_after = Pt(12)

        if data.get('courses'):
            doc.add_heading('Courses', level=1)
            for course in data.get('courses', []):
                p = doc.add_paragraph()
                p.add_run((course.get('title') or '') + '\n').bold = True
                p.add_run(f"{(course.get('institution') or '')}\n")
                p.add_run((course.get('dates') or '') + '\n')
                if course.get('description'):
                    p.add_run(f"- {(course.get('description') or '')}")
                p.paragraph_format.space_after = Pt(12)

        if data.get('projects'):
            doc.add_heading('Projects', level=1)
            for proj in data.get('projects', []):
                p = doc.add_paragraph()
                p.add_run((proj.get('title') or '') + '\n').bold = True
                description_points = (proj.get('description') or '').split('\n')
                formatted_description = "".join([f"- {point.strip()}\n" for point in description_points if point.strip()])
                p.add_run(formatted_description)
                if proj.get('url'):
                    p.add_run(f"URL: {proj.get('url') or ''}")
                p.paragraph_format.space_after = Pt(12)

        if data.get('soft_skills'):
            doc.add_heading('Soft Skills', level=1)
            for skill in data.get('soft_skills', []):
                doc.add_paragraph(skill, style='List Bullet')

        if data.get('skills'):
            doc.add_heading('Skills', level=1)
            skills_list = data.get('skills', [])
            if skills_list:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                for i, skill in enumerate(skills_list):
                    row = i // 3
                    col = i % 3
                    if col == 0 and row > 0:
                        table.add_row()
                    table.cell(row, col).text = skill

        if data.get('languages'):
            doc.add_heading('Languages', level=1)
            for lang in data.get('languages', []):
                doc.add_paragraph(lang, style='List Bullet')

        # --- ALTERAÇÃO FINAL: Salvar em memória ---
        safe_name = secure_filename(data.get('name') or 'Candidato')
        output_filename = f"Curriculo_{safe_name}.docx"

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0) # Retorna o cursor para o início do buffer

        logging.info(f"Documento DOCX '{output_filename}' criado com sucesso em memória.")
        return file_stream, output_filename

    except Exception as e:
        logging.error(f"Exceção ao criar o DOCX em memória: {e}", exc_info=True)
        raise e

# --- 3. ROTAS DA APLICAÇÃO FLASK ---

@app.route('/', methods=['GET', 'POST'])
def handle_all():
    """Esta única rota lida com o GET para a página e o POST para o upload."""
    if request.method == 'POST':
        return upload_file()
    # Se for GET, apenas renderiza a página principal
    return render_template('index.html')

def upload_file():
    logging.info("="*50)
    logging.info("Recebida nova requisição de upload.")

    if 'file' not in request.files:
        return jsonify({'status': 'error', 'message': 'Nenhum arquivo enviado'}), 400

    file = request.files['file']

    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'status': 'error', 'message': 'Nenhum arquivo selecionado ou formato inválido.'}), 400

    try:
        # Lê o arquivo diretamente do stream de upload, sem salvar em disco
        texto_curriculo, sucesso = ler_arquivo_de_stream(file.stream, file.filename)
        if not sucesso:
            return jsonify({'status': 'error', 'message': texto_curriculo}), 500

        logging.info("Texto extraído. Enviando para a API do Google.")
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key:
            logging.critical("CRÍTICO: A chave da API do Google não foi encontrada.")
            return jsonify({'status': 'error', 'message': 'Erro crítico de configuração no servidor.'}), 500

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.5-flash')

        prompt = f"""
        Você é um assistente de RH especialista em analisar currículos. Sua tarefa é extrair as informações do texto de um currículo e retorná-las ESTRITAMENTE em formato JSON.
        A estrutura deve ser:
        {{
          "name": "string",
          "role": "string",
          "contact_info": {{"email": "string", "phone": "string", "location": "string", "linkedin": "string", "github": "string"}},
          "profile": "string",
          "professional_experience": [{{"title": "string", "company": "string", "location": "string", "dates": "string", "description": "string com \\n para cada bullet point"}}, ...],
          "education": [{{"degree": "string", "institution": "string", "dates": "string", "description": "string (opcional)"}}, ...],
          "courses": [{{"title": "string", "institution": "string", "dates": "string", "description": "string (opcional)"}}, ...],
          "projects": [{{"title": "string", "description": "string com \\n para cada bullet point", "url": "string"}}, ...],
          "soft_skills": ["string", "string", ...],
          "skills": ["string", "string", ...],
          "languages": ["string", "string (ex: Inglês (Fluente))", ...]
        }}
        Se uma informação não for encontrada, use um valor vazio (""), null, ou uma lista vazia ([]). Não invente informações.

        TEXTO DO CURRÍCULO PARA ANÁLISE:
        ---
        {texto_curriculo}
        ---
        """

        response = model.generate_content(prompt)
        logging.info("--- RESPOSTA BRUTA DA IA ---")
        logging.info(response.text)

        if not response.text:
             return jsonify({'status': 'error', 'message': 'A IA não conseguiu processar o currículo.'}), 500

        json_text = response.text.strip().replace('```json', '').replace('```', '')
        dados_extraidos = json.loads(json_text)

        if not (dados_extraidos and dados_extraidos.get('name')):
            return jsonify({'status': 'error', 'message': 'A IA não conseguiu extrair informações úteis do documento.'}), 500

        # --- MUDANÇA PRINCIPAL: Gera e retorna o arquivo diretamente ---
        buffer_memoria, nome_arquivo_saida = criar_curriculo_docx_em_memoria(dados_extraidos)

        # Retorna o arquivo gerado em memória como uma resposta de download
        return Response(
            buffer_memoria,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={"Content-Disposition": f"attachment;filename={nome_arquivo_saida}"}
        )

    except json.JSONDecodeError as e:
        logging.error(f"ERRO CRÍTICO: Erro de decodificação do JSON: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': 'Erro ao decodificar a resposta da IA. Não era um JSON válido.'}), 500
    except Exception as e:
        logging.error(f"ERRO CRÍTICO: Uma exceção não tratada ocorreu: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': f'Ocorreu um erro inesperado no processamento: {e}'}), 500