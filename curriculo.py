from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Document
doc = Document()

# Set default font for the document
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Add a Title
title = doc.add_paragraph('Name') 
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title.style = 'Title'
title.runs[0].font.size = Pt(24)
title.runs[0].font.bold = True

# Add a Subtitle (Role)
subtitle = doc.add_paragraph('Full Stack Intern | Computer Science')
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle.style = 'Subtitle'
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)  # Gray color

# Add Contact Information
contact_info = doc.add_paragraph()
contact_info.add_run('São Paulo , Brazil | +55 11 20052004 | email\n')
contact_info.add_run('LinkedIn: URL | ')
contact_info.add_run('GitHub: URL')
contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
contact_info.runs[0].font.size = Pt(11)

# Add a horizontal line
def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.add_run('_' * 80)  # Add a line of underscores
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].font.size = Pt(12)

add_horizontal_line(doc)

# Add Profile Section
doc.add_heading('Profile', level=1)
profile = doc.add_paragraph(
    'Lorem ipsum dolor sit amet, consectetur adipiscing elit. Nullam in tellus nec dolor volutpat tempor sit amet nec purus.'
    'Aenean volutpat neque eget nibh dictum luctus. Nunc erat libero, tristique at orci sagittis, interdum tristique tellus.'
    'Sed placerat lectus non mollis vestibulum. Fusce aliquet turpis ut bibendum mollis.'
)
profile.paragraph_format.space_after = Pt(12)  # Add spacing after the paragraph

# Add Professional Experience Section
doc.add_heading('Professional Experience', level=1)
exp = doc.add_paragraph()
exp.add_run('Full Stack Developer\n').bold = True
exp.add_run('Google, São Paulo\n')
exp.add_run('January 2025 – Present\n')
exp.add_run(
    '- Lorem ipsum dolor sit amet, consectetur adipiscing elit. Nullam in tellus nec dolor volutpat tempor sit amet nec purus.\n'
    '- Nunc erat libero, tristique at orci sagittis, interdum tristique tellus.\n'
    '- Sed placerat lectus non mollis vestibulum. Fusce aliquet turpis ut bibendum mollis.\n'
    '- Quisque dapibus, urna quis ultricies lacinia, nisi sem feugiat quam, at scelerisque massa sapien ac mauris.'
)
exp.paragraph_format.space_after = Pt(12)

# Add Education Section
doc.add_heading('Education', level=1)
edu = doc.add_paragraph()
edu.add_run('Bachelor\'s in Computer Science\n').bold = True
edu.add_run('University of São Paulo (USP)\n')
edu.add_run('February 2025 – December 2028\n')
edu.add_run('- Currently in the second semester with a focus on database management.')
edu.paragraph_format.space_after = Pt(12)

# Add Courses Section
doc.add_heading('Courses', level=1)
courses = doc.add_paragraph()
courses.add_run('Software Engineering\n').bold = True
courses.add_run('Presential Course – São Paulo\n')
courses.add_run('January 2025 – August 2025\n')
courses.add_run('- Attended in-person classes focused on software development, algorithms, and database management.')
courses.paragraph_format.space_after = Pt(12)
""
# Add Projects Section
doc.add_heading('Projects', level=1)
projects = [
{
    "title": "Personal Portfolio Website",
    "description": "A modern and responsive personal portfolio built with HTML, CSS, and JavaScript.\n"
                    "- Designed a clean UI/UX to showcase projects, skills, and contact information.\n"
                    "- Implemented smooth animations and transitions using CSS and vanilla JavaScript.\n"
                    "- Optimized for fast performance and mobile responsiveness with a grid/flexbox layout.\n"
                    "- Integrated a contact form with email service for direct communication.\n"
                    "- Deployed on GitHub Pages with version control and continuous updates.\n",
    "url": "URL"
}
]

for project in projects:
    p = doc.add_paragraph()
    p.add_run(f"{project['title']}\n").bold = True
    p.add_run(f"{project['description']}\n")
    if 'url' in project:  # Check if 'url' key exists
        p.add_run(f"URL: {project['url']}")
    p.paragraph_format.space_after = Pt(12)

# Add Soft Skills Section
doc.add_heading('Soft Skills', level=1)
soft_skills = [
    'Lorem ipsum dolor sit amet, consectetur adipiscing elit',
    'Nullam in tellus nec dolor volutpat tempor sit amet nec purus. ',
    'Aenean volutpat neque eget nibh dictum luctus.',
    'Nunc erat libero, tristique at orci sagittis, interdum tristique tellus.',
    'Sed placerat lectus non mollis vestibulum. Fusce aliquet turpis ut bibendum mollis.'
]
for skill in soft_skills:
    doc.add_paragraph(skill, style='List Bullet')

# Add Skills Section
doc.add_heading('Skills', level=1)
skills = [
    'Python', 'Django', 'Flask', 'FastAPI', 'Java', 'Spring Boot',
    'MongoDB', 'PostgreSQL', 'Docker', 'Kubernetes', 'AWS', 'Figma'
]
skills_table = doc.add_table(rows=1, cols=3)
skills_table.autofit = True
skills_table.style = 'Table Grid'

# Add skills in a 3-column table
for i, skill in enumerate(skills):
    row = i // 3
    col = i % 3
    if col == 0 and row > 0:
        skills_table.add_row()
    skills_table.rows[row].cells[col].text = skill

# Add Languages Section
doc.add_heading('Languages', level=1)
doc.add_paragraph('English (Fluent)', style='List Bullet')

# Save the Document
docx_file = 'Curriculo.docx'
doc.save(docx_file)

print(f"Document created successfully: {docx_file}")