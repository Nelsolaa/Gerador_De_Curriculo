import os
import json
import re
import logging
from flask import Flask, render_template, request, send_from_directory, flash, redirect, url_for
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from flask import Flask, render_template, request, send_from_directory, flash, redirect, url_for, jsonify
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pypdf import PdfReader

# --- 1. CONFIGURAÇÃO INICIAL ---
load_dotenv()

# Configuração de Logs
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("app.log", encoding='utf-8'),
        logging.StreamHandler()
    ]
)

# Constantes e configuração do Flask
UPLOAD_FOLDER = 'uploads'
GENERATED_FOLDER = 'generated'
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['GENERATED_FOLDER'] = GENERATED_FOLDER
app.config['SECRET_KEY'] = os.getenv('SEGREDO_FLASK')

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(GENERATED_FOLDER, exist_ok=True)

# --- 3. FUNÇÕES AUXILIARES ---
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def ler_arquivo(caminho_do_arquivo):
    logging.info(f"Iniciando a leitura do arquivo: {caminho_do_arquivo}")
    if not os.path.exists(caminho_do_arquivo):
        logging.error(f"Arquivo não encontrado em: {caminho_do_arquivo}")
        return "Erro: Arquivo não encontrado.", False
    texto_completo = ""
    try:
        if caminho_do_arquivo.lower().endswith(".pdf"):
            leitor_pdf = PdfReader(caminho_do_arquivo)
            for pagina in leitor_pdf.pages:
                texto_completo += pagina.extract_text() or ""
        elif caminho_do_arquivo.lower().endswith(".docx"):
            doc = Document(caminho_do_arquivo)
            for paragrafo in doc.paragraphs:
                texto_completo += paragrafo.text + "\n"
        logging.info("Arquivo lido com sucesso.")
    except Exception as e:
        logging.error(f"Exceção ao ler o arquivo: {e}", exc_info=True)
        return f"Erro ao ler o arquivo: {e}", False
    return texto_completo, True

def criar_curriculo_docx(data, output_folder):
    """Gera um documento .docx formatado a partir dos dados extraídos pela IA."""
    logging.info("Iniciando a criação do documento DOCX.")
    try:
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Usando 'or' para garantir que valores None se tornem strings vazias
        contact_info_data = data.get('contact_info') or {}
        
        doc.add_paragraph(data.get('name') or 'Nome não encontrado', style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(data.get('role') or 'Cargo não encontrado', style='Subtitle').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        contact_line = " | ".join(filter(None, [
            contact_info_data.get('location'), 
            contact_info_data.get('phone'), 
            contact_info_data.get('email')
        ]))
        doc.add_paragraph(contact_line).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        links_line = " | ".join(filter(None, [
            f"LinkedIn: {contact_info_data.get('linkedin')}" if contact_info_data.get('linkedin') else None,
            f"GitHub: {contact_info_data.get('github')}" if contact_info_data.get('github') else None
        ]))
        doc.add_paragraph(links_line).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph('_' * 80).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if data.get('profile'):
            doc.add_heading('Profile', level=1)
            p = doc.add_paragraph(data.get('profile') or '')
            p.paragraph_format.space_after = Pt(12)

        if data.get('professional_experience'):
            doc.add_heading('Professional Experience', level=1)
            for exp in data.get('professional_experience', []):
                p = doc.add_paragraph()
                p.add_run((exp.get('title') or '') + '\n').bold = True
                p.add_run(f"{(exp.get('company') or '')}, {(exp.get('location') or '')}\n")
                p.add_run((exp.get('dates') or '') + '\n') # <-- CORRIGIDO AQUI
                description_points = (exp.get('description') or '').split('\n')
                formatted_description = "".join([f"- {point.strip()}\n" for point in description_points if point.strip()])
                p.add_run(formatted_description)
                p.paragraph_format.space_after = Pt(12)

        if data.get('education'):
            doc.add_heading('Education', level=1)
            for edu in data.get('education', []):
                p = doc.add_paragraph()
                p.add_run((edu.get('degree') or '') + '\n').bold = True
                p.add_run((edu.get('institution') or '') + '\n')
                p.add_run((edu.get('dates') or '') + '\n')
                if edu.get('description'):
                    p.add_run(f"- {(edu.get('description') or '')}")
                p.paragraph_format.space_after = Pt(12)
        
        # (O resto da função continua igual, pois já era robusta)
        if data.get('courses'):
            doc.add_heading('Courses', level=1)
            for course in data.get('courses', []):
                p = doc.add_paragraph()
                p.add_run((course.get('title') or '') + '\n').bold = True
                p.add_run(f"{(course.get('institution') or '')}\n")
                p.add_run((course.get('dates') or '') + '\n')
                if course.get('description'):
                    p.add_run(f"- {(course.get('description') or '')}")
                p.paragraph_format.space_after = Pt(12)

        if data.get('projects'):
            doc.add_heading('Projects', level=1)
            for proj in data.get('projects', []):
                p = doc.add_paragraph()
                p.add_run((proj.get('title') or '') + '\n').bold = True
                description_points = (proj.get('description') or '').split('\n')
                formatted_description = "".join([f"- {point.strip()}\n" for point in description_points if point.strip()])
                p.add_run(formatted_description)
                p.add_run(f"URL: {proj.get('url') or ''}")
                p.paragraph_format.space_after = Pt(12)

        if data.get('soft_skills'):
            doc.add_heading('Soft Skills', level=1)
            for skill in data.get('soft_skills', []):
                doc.add_paragraph(skill, style='List Bullet')

        if data.get('skills'):
            doc.add_heading('Skills', level=1)
            skills_list = data.get('skills', [])
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            for i, skill in enumerate(skills_list):
                row = i // 3
                col = i % 3
                if col == 0 and row > 0:
                    table.add_row()
                table.cell(row, col).text = skill
        
        if data.get('languages'):
            doc.add_heading('Languages', level=1)
            for lang in data.get('languages', []):
                doc.add_paragraph(lang, style='List Bullet')

        safe_name = secure_filename(data.get('name') or 'Candidato')
        output_filename = f"Curriculo_{safe_name}.docx"
        output_path = os.path.join(output_folder, output_filename)
        doc.save(output_path)
        logging.info(f"Documento DOCX '{output_filename}' criado com sucesso.")
        return output_filename
    except Exception as e:
        logging.error(f"Exceção ao criar o DOCX: {e}", exc_info=True)
        raise e

# --- 4. ROTAS DA APLICAÇÃO FLASK ---

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    logging.info("="*50)
    logging.info("Recebida nova requisição de upload.")
    
    if 'file' not in request.files:
        logging.warning("Requisição recebida sem a parte do arquivo ('file').")
        return jsonify({'status': 'error', 'message': 'Nenhum arquivo enviado'}), 400
    
    file = request.files['file']

    if file.filename == '' or not allowed_file(file.filename):
        logging.warning(f"Upload tentado com nome de arquivo vazio ou tipo não permitido: '{file.filename}'")
        return jsonify({'status': 'error', 'message': 'Nenhum arquivo selecionado ou formato inválido.'}), 400

    try:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        logging.info(f"Arquivo salvo em: '{filepath}'")

        texto_curriculo, sucesso = ler_arquivo(filepath)
        if not sucesso:
            return jsonify({'status': 'error', 'message': texto_curriculo}), 500

        logging.info("Texto extraído. Enviando para a API do Google.")
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key:
            logging.critical("CRÍTICO: A chave da API do Google não foi encontrada.")
            return jsonify({'status': 'error', 'message': 'Erro crítico de configuração no servidor.'}), 500
            
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        # --- PROMPT COMPLETO E CORRIGIDO ---
        prompt = f"""
        Você é um assistente de RH especialista em analisar currículos. Sua tarefa é extrair as informações do texto de um currículo e retorná-las ESTRITAMENTE em formato JSON.
        A estrutura deve ser:
        {{
          "name": "string",
          "role": "string",
          "contact_info": {{"email": "string", "phone": "string", "location": "string", "linkedin": "string", "github": "string"}},
          "profile": "string",
          "professional_experience": [{{"title": "string", "company": "string", "location": "string", "dates": "string", "description": "string com \\n para cada bullet point"}}, ...],
          "education": [{{"degree": "string", "institution": "string", "dates": "string", "description": "string (opcional)"}}, ...],
          "courses": [{{"title": "string", "institution": "string", "dates": "string", "description": "string (opcional)"}}, ...],
          "projects": [{{"title": "string", "description": "string com \\n para cada bullet point", "url": "string"}}, ...],
          "soft_skills": ["string", "string", ...],
          "skills": ["string", "string", ...],
          "languages": ["string", "string (ex: Inglês (Fluente))", ...]
        }}
        Se uma informação não for encontrada, use um valor vazio (""), null, ou uma lista vazia ([]). Não invente informações.

        TEXTO DO CURRÍCULO PARA ANÁLISE:
        ---
        {texto_curriculo}
        ---
        """
        
        safety_config = {
            HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_ONLY_HIGH,
            HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_ONLY_HIGH,
            HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_ONLY_HIGH,
            HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_ONLY_HIGH,
        }

        response = model.generate_content(prompt, safety_settings=safety_config)
        
        logging.info("--- RESPOSTA BRUTA DA IA ---")
        logging.info(response.text)

        if not response.text:
            logging.warning(f"Resposta da IA vazia. Feedback de segurança: {response.prompt_feedback}")
            return jsonify({'status': 'error', 'message': 'A IA não conseguiu processar o currículo devido aos filtros de segurança.'}), 500

        json_text = response.text.strip().replace('```json', '').replace('```', '')
        dados_extraidos = json.loads(json_text)
        
        if not (dados_extraidos and dados_extraidos.get('name')):
            logging.warning("ALERTA: A IA retornou um JSON válido, mas ele parece estar VAZIO.")
            return jsonify({'status': 'error', 'message': 'A IA não conseguiu extrair informações úteis do documento.'}), 500

        nome_arquivo_saida = criar_curriculo_docx(dados_extraidos, app.config['GENERATED_FOLDER'])
        
        # SUCESSO: Retorna o link para download em formato JSON
        download_url = url_for('download_file', filename=nome_arquivo_saida)
        return jsonify({'status': 'success', 'download_url': download_url})

    except json.JSONDecodeError as e:
        logging.error(f"ERRO CRÍTICO: Erro de decodificação do JSON: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': 'Erro ao decodificar a resposta da IA. Não era um JSON válido.'}), 500
    except Exception as e:
        logging.error(f"ERRO CRÍTICO: Uma exceção não tratada ocorreu: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': f'Ocorreu um erro inesperado no processamento: {e}'}), 500

@app.route('/download/<filename>')
def download_file(filename):
    logging.info(f"Iniciando download do arquivo: {filename}")
    return send_from_directory(app.config['GENERATED_FOLDER'], filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)